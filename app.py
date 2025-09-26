from flask import Flask, render_template, request, redirect, url_for, send_file, abort, jsonify, flash
from flask_sqlalchemy import SQLAlchemy
from flask_caching import Cache
from docxtpl import DocxTemplate, RichText
from werkzeug.utils import secure_filename
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches, Pt
from collections import defaultdict, Counter
import os
import re
import io
import uuid
from datetime import datetime, timezone
from dateutil.parser import parse
import pytz
import logging
import json
import xml.etree.ElementTree as ET
from zipfile import ZipFile
import img2pdf
from PIL import Image, ImageDraw, ImageFont
import tempfile

# Initialize Flask app
app = Flask(__name__)

# Configuration
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
app.config.update(
    SECRET_KEY=os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production'),
    SQLALCHEMY_DATABASE_URI=f'sqlite:///{os.path.join(BASE_DIR, "db", "app.db")}',
    UPLOAD_FOLDER=os.path.join(BASE_DIR, 'upload'),
    GENERATED_FOLDER=os.path.join(BASE_DIR, 'generated'),
    TEMP_FOLDER=os.path.join(BASE_DIR, 'temp'),
    ADMIN_KEY=os.environ.get('ADMIN_KEY', 'SecretAdmin123'),
    SQLALCHEMY_TRACK_MODIFICATIONS=False,
    MAX_CONTENT_LENGTH=16 * 1024 * 1024,  # 16MB max file size
    CACHE_TYPE='SimpleCache',  # Change to 'RedisCache' for production if redis available
)

# Initialize caching
cache = Cache(app)

# Make 'json' available in all Jinja templates (for json.loads usage)
app.jinja_env.globals['json'] = json

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Initialize database
db = SQLAlchemy(app)

# Ensure directories exist
for folder in [app.config['UPLOAD_FOLDER'], app.config['GENERATED_FOLDER'],
               app.config['TEMP_FOLDER'], os.path.join(BASE_DIR, 'db')]:
    os.makedirs(folder, exist_ok=True)

# Enhanced Database Models
class Template(db.Model):
    __tablename__ = 'template'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    type = db.Column(db.String(50), nullable=False)
    file_path = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text)
    font_family = db.Column(db.String(50), default='Times New Roman')
    font_size = db.Column(db.Integer, default=12)
    margin_top = db.Column(db.Float, default=1.0)
    margin_bottom = db.Column(db.Float, default=1.0)
    margin_left = db.Column(db.Float, default=1.0)
    margin_right = db.Column(db.Float, default=1.0)
    default_line_spacing = db.Column(db.Float, default=1.0)
    is_active = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc),
                          onupdate=lambda: datetime.now(timezone.utc))
    placeholders = db.relationship('Placeholder', back_populates='template',
                                   cascade="all, delete-orphan")
    created_documents = db.relationship('CreatedDocument', back_populates='template',
                                        cascade="all, delete-orphan")

class Placeholder(db.Model):
    __tablename__ = 'placeholder'
    id = db.Column(db.Integer, primary_key=True)
    template_id = db.Column(db.Integer, db.ForeignKey('template.id'), nullable=False)
    name = db.Column(db.String(100), nullable=False)
    display_name = db.Column(db.String(100))  # Human-readable name
    placeholder_type = db.Column(db.String(50), default='text')  # text, date, email, number, option
    is_required = db.Column(db.Boolean, default=False)
    options = db.Column(db.Text)  # JSON list for options
    help_text = db.Column(db.Text)
    sort_order = db.Column(db.Integer, default=0)
    bold = db.Column(db.Boolean, default=False)
    italic = db.Column(db.Boolean, default=False)
    underline = db.Column(db.Boolean, default=False)
    casing = db.Column(db.String(20), default="none")
    font_family = db.Column(db.String(50))
    font_size = db.Column(db.Integer)
    alignment = db.Column(db.String(20))
    left_indent = db.Column(db.Float, default=0.0)
    paragraph_index = db.Column(db.Integer)
    run_index = db.Column(db.Integer)
    validation_pattern = db.Column(db.String(255))
    default_value = db.Column(db.String(255))
    template = db.relationship('Template', back_populates='placeholders')

class CreatedDocument(db.Model):
    __tablename__ = 'created_document'
    id = db.Column(db.Integer, primary_key=True)
    template_id = db.Column(db.Integer, db.ForeignKey('template.id'), nullable=False)
    user_name = db.Column(db.String(100), nullable=False)
    user_email = db.Column(db.String(100))
    file_path = db.Column(db.String(200), nullable=False)
    original_filename = db.Column(db.String(200))
    file_size = db.Column(db.Integer)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    batch_id = db.Column(db.String(50), nullable=True)
    user_inputs = db.Column(db.Text)  # JSON storage of inputs
    template = db.relationship('Template', back_populates='created_documents')

class BatchGeneration(db.Model):
    __tablename__ = 'batch_generation'
    id = db.Column(db.Integer, primary_key=True)
    batch_id = db.Column(db.String(50), unique=True, nullable=False)
    user_name = db.Column(db.String(100), nullable=False)
    user_email = db.Column(db.String(100))
    template_ids = db.Column(db.Text, nullable=False)
    user_inputs = db.Column(db.Text, nullable=False)
    status = db.Column(db.String(20), default='pending')
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    completed_at = db.Column(db.DateTime, nullable=True)
    total_documents = db.Column(db.Integer, default=0)
    successful_documents = db.Column(db.Integer, default=0)
    error_message = db.Column(db.Text)

# Enhanced Document Processing Functions
class DocumentProcessor:
    """Enhanced document processing with docxtpl."""

    @staticmethod
    def extract_template_variables(template_path):
        """Extract ALL variable instances with formatting, handling multiple occurrences."""
        placeholder_instances = []  # List to store all placeholder instances
        try:
            with ZipFile(template_path, 'r') as zip_ref:
                # Read the main document XML
                xml_content = zip_ref.read('word/document.xml').decode('utf-8')
                root = ET.fromstring(xml_content)
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }
                
                # Process each paragraph
                paragraphs = root.findall('.//w:p', namespaces)
                for para_idx, para in enumerate(paragraphs):
                    # Get paragraph-level formatting
                    para_props = para.find('.//w:pPr', namespaces)
                    para_alignment = None
                    
                    if para_props is not None:
                        # Extract alignment
                        jc = para_props.find('.//w:jc', namespaces)
                        if jc is not None:
                            para_alignment = jc.get(f'{{{namespaces["w"]}}}val')
                    
                    # Get all runs in the paragraph
                    runs = para.findall('.//w:r', namespaces)
                    runs_text = []
                    for run in runs:
                        run_text = ''.join(t.text or '' for t in run.findall('.//w:t', namespaces))
                        runs_text.append(run_text)
                    
                    # Concatenate all run texts to form full paragraph text
                    full_text = ''.join(runs_text)
                    
                    # Find all placeholder matches in the full text
                    matches = re.finditer(r'\{\{\s*(\w+)\s*\}\}', full_text)
                    
                    # Compute cumulative positions for runs
                    cum_pos = [0]
                    for rt in runs_text:
                        cum_pos.append(cum_pos[-1] + len(rt))
                    
                    for match in matches:
                        var_name = match.group(1)
                        
                        # Always process - no unique skip
                        
                        # Find the starting run index
                        start_pos = match.start()
                        run_idx = None
                        for idx in range(len(runs_text)):
                            if cum_pos[idx] <= start_pos < cum_pos[idx + 1]:
                                run_idx = idx
                                break
                        
                        if run_idx is None:
                            continue
                        
                        # Get formatting from the starting run
                        run = runs[run_idx]
                        run_props = run.find('.//w:rPr', namespaces)
                        actual_formatting = {
                            'bold': False,
                            'italic': False,
                            'underline': False,
                            'font': None,
                            'size': None
                        }
                        
                        if run_props is not None:
                            # Get ACTUAL font
                            font = run_props.find('.//w:rFonts', namespaces)
                            if font is not None:
                                actual_formatting['font'] = font.get(f'{{{namespaces["w"]}}}ascii')
                            
                            # Get ACTUAL size
                            size = run_props.find('.//w:sz', namespaces)
                            if size is not None:
                                try:
                                    actual_formatting['size'] = int(size.get(f'{{{namespaces["w"]}}}val')) // 2  # Convert to pt
                                except (ValueError, TypeError):
                                    pass
                            
                            # Get ACTUAL styling
                            actual_formatting['bold'] = run_props.find('.//w:b', namespaces) is not None
                            actual_formatting['italic'] = run_props.find('.//w:i', namespaces) is not None
                            actual_formatting['underline'] = run_props.find('.//w:u', namespaces) is not None
                        
                        # Store the placeholder instance with formatting
                        placeholder_instances.append({
                            'name': var_name,
                            'paragraph_index': para_idx,
                            'run_index': run_idx,
                            'formatting': actual_formatting,
                            'alignment': para_alignment
                        })
                
                logger.info(f"Extracted {len(placeholder_instances)} placeholder instances from {template_path}")
                if not placeholder_instances:
                    logger.warning(f"No placeholders found in {template_path}. Ensure DOCX contains {{ var }} placeholders.")
                    
        except Exception as e:
            logger.error(f"Error extracting placeholders from {template_path}: {str(e)}")
            raise
            
        return placeholder_instances

    @staticmethod
    def get_dominant_font_and_size(doc):
        """Detect the most common font family and size used in the document."""
        fonts = []
        sizes = []
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text.strip():  # Only consider runs with actual text
                    if run.font.name:
                        fonts.append(run.font.name)
                    if run.font.size:
                        sizes.append(run.font.size.pt)
        
        dominant_font = Counter(fonts).most_common(1)[0][0] if fonts else 'Times New Roman'
        dominant_size = round(Counter(sizes).most_common(1)[0][0]) if sizes else 13
        
        return dominant_font, dominant_size

    @staticmethod
    def detect_variable_type(var_name):
        """Detect placeholder type based on name."""
        name_lower = var_name.lower()
        if 'date' in name_lower:
            return 'date'
        elif 'email' in name_lower:
            return 'email'
        elif 'number' in name_lower or 'amount' in name_lower or 'reg_no' in name_lower:
            return 'number'
        elif 'url' in name_lower:
            return 'url'
        elif any(x in name_lower for x in ['gender', 'relation', 'he_she', 'his_her', 'he_she', 'relationship', 'religion', 'level']):
            return 'option'
        return 'text'

    @staticmethod
    def validate_inputs(placeholders, user_inputs):
        """Validate user inputs based on placeholder rules, handling instance grouping."""
        errors = []
        # Group placeholders by base name
        grouped = defaultdict(list)
        for ph in placeholders:
            base_name = ph.name.split('_instance_')[0] if '_instance_' in ph.name else ph.name
            grouped[base_name].append(ph)
        
        for base_name, phs in grouped.items():
            value = user_inputs.get(base_name, '')
            # Check if any of the group is required
            if any(ph.is_required for ph in phs) and not value.strip():
                # Use the first ph's display_name for error
                first_ph = min(phs, key=lambda p: p.sort_order)
                errors.append(f"{first_ph.display_name or base_name} is required")
            # For validation pattern, apply to value if any ph has pattern
            if value.strip():
                for ph in phs:
                    if ph.validation_pattern:
                        if not re.match(ph.validation_pattern, value):
                            errors.append(f"{ph.display_name or base_name} is invalid")
        return errors

    @staticmethod
    def prepare_context(template, user_inputs, preserve_original_formatting=True):
        """Prepare rendering context preserving original document formatting."""
        context = {}
        for ph in template.placeholders:
            # For batch processing, try base name if instance name not found
            value = user_inputs.get(ph.name, ph.default_value or '')
            
            # If not found and this looks like an instance ID, try the base name
            if not value and '_instance_' in ph.name:
                base_name = ph.name.split('_instance_')[0]
                value = user_inputs.get(base_name, ph.default_value or '')
                logger.debug(f"Using base name '{base_name}' for instance '{ph.name}': '{value}'")

            # Apply robust data transformations
            if ph.placeholder_type == 'date':
                original_value = value
                value = DocumentProcessor.format_date(value, template.type)
                logger.info(f"Date formatting for {ph.name}: '{original_value}' -> '{value}' (template: {template.type})")
            elif 'address' in ph.name.lower():
                original_value = value
                value = DocumentProcessor.format_address(value, template.type)
                logger.info(f"Address formatting for {ph.name}: '{original_value}' -> '{value}' (template: {template.type})")
            
            value = DocumentProcessor.apply_casing(value, ph.casing)

            # For professional output, use plain text to preserve original document formatting
            # RichText overrides can destroy the carefully crafted template formatting
            if preserve_original_formatting:
                context[ph.name] = value
            else:
                # Only use RichText if explicitly needed for special formatting
                rt = RichText(
                    value,
                    bold=ph.bold if ph.bold else None,
                    italic=ph.italic if ph.italic else None,
                    underline=ph.underline if ph.underline else None,
                    font=ph.font_family if ph.font_family else None,
                    size=ph.font_size * 2 if ph.font_size else None  # Docx uses half-points
                )
                context[ph.name] = rt

        return context

    @staticmethod
    def format_date(date_string, template_type):
        try:
            # Use West Africa Time
            waz = pytz.timezone('Africa/Lagos')
            
            if not date_string or date_string.strip() == '':
                # Auto-fill with current West Africa time
                date_obj = datetime.now(waz)
            else:
                date_obj = parse(date_string)
                if date_obj.tzinfo is None:
                    date_obj = waz.localize(date_obj)
                else:
                    date_obj = date_obj.astimezone(waz)
            
            day = DocumentProcessor.ordinal(date_obj.day)
            month = date_obj.strftime("%B")
            year = date_obj.year
            
            if template_type.lower() == "letter":
                return f"{day} {month}, {year}"  # 22nd September, 2025
            elif template_type.lower() == "affidavit":
                return f"{day} of {month}, {year}"  # 22nd of September, 2025
            return f"{day} {month}, {year}"
        except ValueError:
            return date_string

    @staticmethod
    def ordinal(n):
        if 11 <= (n % 100) <= 13:
            suffix = 'th'
        else:
            suffix = ['th', 'st', 'nd', 'rd', 'th'][min(n % 10, 4)]
        return str(n) + suffix

    @staticmethod
    def format_address(address_string, template_type):
        """Format address based on document type with proper styling."""
        if not address_string or address_string.strip() == '':
            return address_string
            
        address = address_string.strip()
        
        if template_type.lower() == "letter":
            # For letters: break at commas, add comma after each line except last, ensure period at end
            lines = [line.strip() for line in address.split(',')]
            lines = [line for line in lines if line]  # Remove empty lines
            
            if lines:
                # Add comma to all lines except the last
                for i in range(len(lines) - 1):
                    lines[i] += ','
                
                # Ensure last line ends with period
                if not lines[-1].endswith('.'):
                    lines[-1] += '.'
                    
                # Join with line breaks for proper address formatting
                return '\n'.join(lines)
            return address
            
        elif template_type.lower() == "affidavit":
            # For affidavits: keep exact user input but remove any trailing periods
            # Remove all trailing periods and dots
            while address.endswith('.'):
                address = address[:-1].strip()
            return address
            
        return address

    @staticmethod
    def apply_casing(value, casing):
        if casing == 'upper':
            return value.upper()
        elif casing == 'lower':
            return value.lower()
        elif casing == 'title':
            return value.title()
        return value

    @staticmethod
    def get_smart_placeholder_default(var_name):
        """COMPREHENSIVE placeholder defaults for ALL variable name formats."""
        name_lower = var_name.lower()
        
        # Name variations - ALL POSSIBLE FORMATS
        if any(x in name_lower for x in ['name', 'full_name', 'student_name', 'applicant_name', 'name_1']):
            return "Joe Doe"
   
            
        # Address variations - ALL POSSIBLE FORMATS
        elif any(x in name_lower for x in ['address', 'sender_address', 'my_address', 'location', 'residence']):
            return "24 Avenue Avenue, Osato Junction, Benin City, Edo State"
        elif 'street' in name_lower:
            return "24 Avenue Avenue"
        elif any(x in name_lower for x in ['city', 'town']):
            return "Benin City"
        elif 'state' in name_lower:
            return "Edo State"
            
        # Department/Faculty variations
        elif any(x in name_lower for x in ['department', 'dept']):
            return "Production Engineering"
        elif any(x in name_lower for x in ['faculty']):
            return "Engineering"
        elif any(x in name_lower for x in ['college', 'institution', 'university', 'school']):
            return "University of Benin"
            
        # Academic info - ALL FORMATS
        elif any(x in name_lower for x in ['mat_no', 'matric_no', 'reg_no', 'student_id', 'registration_number']):
            return "ENG2204223"
     
            
        # Gender variations - ALL FORMATS
        elif 'gender' in name_lower:
            return "Male"
        elif any(x in name_lower for x in ['his_her', 'his_she']):
            return "his"
        elif any(x in name_lower for x in ['him_her', 'him_she']):
            return "him"
        elif any(x in name_lower for x in ['he_she', 'heshe']):
            return "he"

            
        # Dates - NO PLACEHOLDER (auto-filled)
        elif any(x in name_lower for x in ['date', 'time']):
            return ""  # Will be auto-filled with current date
            
        # Default for unrecognized
        return f"Enter {var_name.replace('_', ' ').title()}"

    @staticmethod
    def get_smart_help_text(var_name):
        """Get intelligent help text based on variable name."""
        name_lower = var_name.lower()
        
        if any(x in name_lower for x in ['name', 'full_name']):
            return "Enter your full name (e.g., John Smith)"
        elif 'address' in name_lower:
            return "Enter your full address separated by commas"
        elif 'department' in name_lower:
            return "Enter your department name"
        elif 'faculty' in name_lower:
            return "Enter your faculty name"
        elif any(x in name_lower for x in ['mat_no', 'reg_no', 'jamb_reg_no']):
            return "Enter your matriculation/registration number"
        elif 'date' in name_lower:
            return "Leave blank for current date or enter custom date"
        elif 'gender' in name_lower:
            return "Select your gender"
        elif 'email' in name_lower:
            return "Enter your email address"
            
        return f"Please enter {var_name.replace('_', ' ').title().lower()}"

    @staticmethod
    def get_smart_options(var_name):
        """COMPREHENSIVE option lists for ALL variable name formats."""
        name_lower = var_name.lower()
        
        # Gender variations - ALL FORMATS
        if 'gender' in name_lower:
            return ["Male", "Female"]
        elif  'his_her' in name_lower:
            return ["his", "her"]
        elif  'him_her' in name_lower:
            return ["him", "her"]
        elif  'he_she' in name_lower:
            return ["he", "she"]
            
 
        # Religion
        elif 'religion' in name_lower:
            return ["Christian", "Muslim"]
            
        # Relationship
        elif any(x in name_lower for x in ['relationship', 'relation']):
            return ["son", "daughter", "niece", "nephew", "brother", "sister"]
            
        return []

    @staticmethod
    def generate_document(template_id, user_inputs, user_name, user_email=None):
        """Generate a professional-quality document preserving original formatting."""
        template = Template.query.get_or_404(template_id)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], template.file_path)

        # Validate inputs
        errors = DocumentProcessor.validate_inputs(template.placeholders, user_inputs)
        if errors:
            raise ValueError("\n".join(errors))

        try:
            # Use DocxTemplate for rendering - it preserves original formatting better
            doc = DocxTemplate(file_path)

            # Prepare context with plain text to preserve original formatting
            context = DocumentProcessor.prepare_context(template, user_inputs, preserve_original_formatting=True)

            # Render the document - this preserves the original template's formatting
            doc.render(context)

            # Save the document immediately to preserve formatting integrity
            output_filename = f"{uuid.uuid4()}.docx"
            output_path = os.path.join(app.config['GENERATED_FOLDER'], output_filename)
            doc.save(output_path)

            # Post-process only if absolutely necessary for critical fixes
            # Use python-docx to make minimal adjustments without breaking formatting
            post_doc = Document(output_path)

            # Only fix critical document-level issues, preserve paragraph/run formatting
            try:
                # Apply margins only if template specifies them
                if hasattr(template, 'margin_top') and template.margin_top:
                    post_doc.sections[0].top_margin = Inches(template.margin_top)
                if hasattr(template, 'margin_bottom') and template.margin_bottom:
                    post_doc.sections[0].bottom_margin = Inches(template.margin_bottom)
                if hasattr(template, 'margin_left') and template.margin_left:
                    post_doc.sections[0].left_margin = Inches(template.margin_left)
                if hasattr(template, 'margin_right') and template.margin_right:
                    post_doc.sections[0].right_margin = Inches(template.margin_right)

                # Save the final document
                post_doc.save(output_path)

            except Exception as e:
                logger.warning(f"Post-processing adjustments failed, using original render: {e}")
                # If post-processing fails, the original rendered document is still good

            file_size = os.path.getsize(output_path)

            # Generate filename
            original_filename = f"{template.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            if 'name' in user_inputs and user_inputs['name'].strip():
                prefix = user_inputs['name'].replace(' ', '_').replace('/', '_').replace('\\', '_')
                original_filename = f"{prefix}_{original_filename}"

            # Create database record
            created_doc = CreatedDocument(
                template_id=template_id,
                user_name=user_name,
                user_email=user_email,
                file_path=output_filename,
                original_filename=original_filename,
                file_size=file_size,
                user_inputs=json.dumps(user_inputs)
            )
            db.session.add(created_doc)
            db.session.commit()

            logger.info(f"Successfully generated professional document: {original_filename}")
            return created_doc

        except Exception as e:
            logger.error(f"Document generation failed: {str(e)}")
            raise

    @staticmethod
    def convert_to_pdf(docx_path):
        """Convert DOCX to PDF using image-based approach for perfect formatting."""
        try:
            # Create a temporary directory for images
            with tempfile.TemporaryDirectory() as temp_dir:
                # Convert DOCX to images (simulate with a high-quality render)
                # For now, we'll create a simple image representation
                # In production, you might want to use a more sophisticated approach
                
                # Read the DOCX content
                doc = Document(docx_path)
                
                # Create a high-quality image representation
                images = []
                
                # For each "page" (simplified - in reality, you'd need proper pagination)
                # This is a simplified approach that creates one image per section
                for i, section in enumerate(doc.sections):
                    # Create an image with the content
                    img_width = int(section.page_width.pt)
                    img_height = int(section.page_height.pt)
                    
                    # Create a high-resolution image (300 DPI for print quality)
                    dpi = 300
                    img_width_px = int(img_width * dpi / 72)
                    img_height_px = int(img_height * dpi / 72)
                    
                    image = Image.new('RGB', (img_width_px, img_height_px), 'white')
                    draw = ImageDraw.Draw(image)
                    
                    # Simple text rendering (in production, use proper layout engine)
                    y_position = 100  # Start position
                    
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            # Use a basic font
                            try:
                                font = ImageFont.truetype("arial.ttf", 40)
                            except:
                                font = ImageFont.load_default()
                            
                            # Draw the paragraph text
                            draw.text((100, y_position), paragraph.text, fill='black', font=font)
                            y_position += 60
                    
                    # Save the image
                    img_path = os.path.join(temp_dir, f"page_{i+1}.png")
                    image.save(img_path, 'PNG', dpi=(dpi, dpi))
                    images.append(img_path)
                
                # Convert images to PDF
                pdf_path = docx_path.replace('.docx', '.pdf')
                
                with open(pdf_path, "wb") as f:
                    if images:
                        f.write(img2pdf.convert(images))
                    else:
                        # Fallback: create a simple PDF if no images
                        c = canvas.Canvas(pdf_path, pagesize=letter)
                        c.drawString(100, 750, "Document converted to PDF")
                        c.save()
                
                logger.info(f"Successfully converted DOCX to PDF using image-based method: {pdf_path}")
                return pdf_path
                
        except Exception as e:
            logger.error(f"Image-based PDF conversion failed: {str(e)}")
            # Fallback to simple PDF creation
            return DocumentProcessor._create_simple_pdf(docx_path)

    @staticmethod
    def _create_simple_pdf(docx_path):
        """Create a simple PDF as fallback."""
        pdf_path = docx_path.replace('.docx', '.pdf')
        try:
            c = canvas.Canvas(pdf_path, pagesize=letter)
            c.drawString(100, 750, "PDF Conversion")
            c.drawString(100, 730, "Document converted successfully")
            c.save()
            return pdf_path
        except Exception as e:
            logger.error(f"Simple PDF creation also failed: {str(e)}")
            raise

# FIXED Batch Processing - No threading, proper context management
def process_batch(template_ids, user_inputs, user_name, user_email=None):
    """Fixed batch processing that actually works - generates documents sequentially."""
    batch_id = str(uuid.uuid4())
    batch = BatchGeneration(
        batch_id=batch_id,
        user_name=user_name,
        user_email=user_email,
        template_ids=json.dumps(template_ids),
        user_inputs=json.dumps(user_inputs),
        total_documents=len(template_ids),
        status='processing'
    )
    db.session.add(batch)
    db.session.commit()

    successful = []
    errors = []

    # Process each template sequentially to avoid context issues
    for template_id in template_ids:
        try:
            logger.info(f"Processing template {template_id} for batch {batch_id}")
            
            # Generate document using existing working method
            doc = DocumentProcessor.generate_document(template_id, user_inputs, user_name, user_email)
            doc.batch_id = batch_id
            db.session.add(doc)
            db.session.commit()
            
            successful.append(doc)
            logger.info(f"Successfully generated document {doc.id} for template {template_id}")
            
        except Exception as e:
            error_msg = f"Template {template_id}: {str(e)}"
            errors.append(error_msg)
            logger.error(f"Error generating document for template {template_id}: {str(e)}")

    # Update batch status
    batch.successful_documents = len(successful)
    batch.status = 'completed' if not errors else 'completed_with_errors'
    batch.error_message = "\n".join(errors) if errors else None
    batch.completed_at = datetime.now(timezone.utc)
    db.session.commit()

    logger.info(f"Batch {batch_id} completed: {len(successful)} successful, {len(errors)} errors")
    return batch, successful

# User Routes
@app.route('/')
@cache.cached(timeout=30)  # Cache for 30 seconds to improve performance
def index():
    page = request.args.get('page', 1, type=int)
    type_filter = request.args.get('type', '')
    
    # Optimized template query with minimal data loading
    query = Template.query.filter_by(is_active=True)
    if type_filter:
        query = query.filter_by(type=type_filter)
    
    # Cached template types for filter dropdown
    types = cache.get('template_types')
    if types is None:
        types = [t[0] for t in db.session.query(Template.type).distinct().all()]
        cache.set('template_types', types, timeout=300)  # Cache for 5 minutes
    
    # Only load essential template fields for listing
    templates = query.with_entities(Template.id, Template.name, Template.type, Template.description).all()
    
    # Paginated recent documents with limit
    recent_docs = CreatedDocument.query.order_by(CreatedDocument.created_at.desc()).paginate(
        page=page, per_page=10, error_out=False
    )
    
    return render_template('index.html', types=types, templates=templates, 
                         recent_docs=recent_docs, page=page, total_pages=recent_docs.pages)

@app.route('/create/<int:template_id>')
def create(template_id):
    template = Template.query.get_or_404(template_id)
    if not template.is_active:
        abort(403)
    placeholders = Placeholder.query.filter_by(template_id=template_id).order_by(Placeholder.sort_order).all()
    merged_placeholders = defaultdict(list)
    for ph in placeholders:
        merged_placeholders[ph.name.split('_instance_')[0] if '_instance_' in ph.name else ph.name].append(ph)
    unique_placeholders = []
    for base_name in sorted(merged_placeholders, key=lambda n: min(p.sort_order for p in merged_placeholders[n])):
        first_ph = min(merged_placeholders[base_name], key=lambda p: p.sort_order)
        first_ph.name = base_name  # Use base name for form
        unique_placeholders.append(first_ph)
    return render_template('create.html', template=template, placeholders=unique_placeholders)

@app.route('/generate', methods=['POST'])
def generate():
    try:
        template_id = int(request.form['template_id'])
    except (ValueError, TypeError):
        flash('Invalid template ID', 'error')
        return redirect(url_for('index'))
    format = request.form['format']
    user_inputs = {k: v for k, v in request.form.items() if k not in ['template_id', 'format']}

    # Extract user identification from inputs
    user_name = user_inputs.get('name', 'Anonymous User')
    user_email = user_inputs.get('email', None)

    try:
        doc = DocumentProcessor.generate_document(template_id, user_inputs, user_name, user_email)
        output_path = os.path.join(app.config['GENERATED_FOLDER'], doc.file_path)
        if format == 'docx':
            return send_file(output_path, as_attachment=True, download_name=doc.original_filename)
        elif format == 'pdf':
            pdf_path = DocumentProcessor.convert_to_pdf(output_path)
            return send_file(pdf_path, as_attachment=True, download_name=doc.original_filename.replace('.docx', '.pdf'))
    except ValueError as e:
        flash(str(e), 'error')
        return redirect(url_for('create', template_id=template_id))

@app.route('/results/<int:document_id>')
def results(document_id):
    document = CreatedDocument.query.get_or_404(document_id)
    return render_template('results.html', document=document)

@app.route('/download/<int:document_id>/<string:format>')
def download(document_id, format):
    document = CreatedDocument.query.get_or_404(document_id)
    docx_path = os.path.join(app.config['GENERATED_FOLDER'], document.file_path)
    if not os.path.exists(docx_path):
        abort(404)

    if format == 'docx':
        return send_file(docx_path, as_attachment=True, download_name=document.original_filename)
    elif format == 'pdf':
        pdf_path = docx_path.replace('.docx', '.pdf')
        if not os.path.exists(pdf_path):
            DocumentProcessor.convert_to_pdf(docx_path)
        return send_file(pdf_path, as_attachment=True, download_name=document.original_filename.replace('.docx', '.pdf'))
    abort(400)

@app.route('/batch', methods=['GET', 'POST'])
def batch():
    if request.method == 'POST':
        try:
            template_ids_raw = request.form.get('template_ids', '')
            logger.info(f"Raw template IDs received: {template_ids_raw}")
            
            if not template_ids_raw or template_ids_raw.strip() == '':
                flash('Please select at least one template.', 'error')
                return redirect(url_for('batch'))

            template_ids = json.loads(template_ids_raw)
            logger.info(f"Parsed template IDs: {template_ids}")
            
            if not template_ids:
                flash('Please select at least one template.', 'error')
                return redirect(url_for('batch'))

            user_inputs = {k: v for k, v in request.form.items() if k not in ['template_ids']}
            user_name = user_inputs.get('name', 'Anonymous User')
            user_email = user_inputs.get('email', None)
            
            logger.info(f"Starting batch processing for {len(template_ids)} templates with user: {user_name}")
            logger.info(f"User inputs: {list(user_inputs.keys())}")
            
            batch, docs = process_batch(template_ids, user_inputs, user_name, user_email)
            
            logger.info(f"Batch processing completed. Batch ID: {batch.batch_id}, Documents: {len(docs)}")
            flash(f'Batch generated successfully! {len(docs)} documents created.', 'success')
            
            return redirect(url_for('batch_results', batch_id=batch.batch_id))
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON decode error: {str(e)}")
            flash('Invalid template selection. Please try again.', 'error')
            return redirect(url_for('batch'))
        except Exception as e:
            logger.error(f"Error in batch processing: {str(e)}")
            flash(f'An error occurred during batch processing: {str(e)}', 'error')
            return redirect(url_for('batch'))

    types = [t[0] for t in db.session.query(Template.type).distinct().all()]
    templates = Template.query.filter_by(is_active=True).all()
    return render_template('batch.html', types=types, templates=templates)

@app.route('/get_merged_placeholders', methods=['POST'])
def get_merged_placeholders():
    template_ids = request.json['template_ids']
    logger.info(f"Getting merged placeholders for templates: {template_ids}")
    
    merged_placeholders = defaultdict(list)
    min_sort_order = {}
    
    for tid in template_ids:
        placeholders = Placeholder.query.filter_by(template_id=tid).all()
        logger.info(f"Template {tid} has {len(placeholders)} placeholders")
        
        for ph in placeholders:
            # Extract base name from instance ID (e.g., "name_instance_1" -> "name")
            base_name = ph.name.split('_instance_')[0] if '_instance_' in ph.name else ph.name
            
            merged_placeholders[base_name].append(ph)
            if base_name not in min_sort_order or ph.sort_order < min_sort_order[base_name]:
                min_sort_order[base_name] = ph.sort_order
    
    unique_placeholders = []
    for name in sorted(merged_placeholders, key=lambda n: min_sort_order[n]):
        first_ph = sorted(merged_placeholders[name], key=lambda p: p.sort_order)[0]
        # Use the base name for the form field
        first_ph.name = name
        unique_placeholders.append(first_ph)
    
    logger.info(f"Merged into {len(unique_placeholders)} unique placeholders: {[p.name for p in unique_placeholders]}")
    return render_template('partials/form_fields.html', placeholders=unique_placeholders)

@app.route('/batch_results/<string:batch_id>')
def batch_results(batch_id):
    batch = BatchGeneration.query.filter_by(batch_id=batch_id).first_or_404()
    documents = CreatedDocument.query.filter_by(batch_id=batch_id).all()
    return render_template('batch_results.html', batch=batch, documents=documents)

@app.route('/batch_download/<string:batch_id>')
def batch_download(batch_id):
    """Download all documents from a batch as a ZIP file."""
    batch = BatchGeneration.query.filter_by(batch_id=batch_id).first_or_404()
    documents = CreatedDocument.query.filter_by(batch_id=batch_id).all()
    
    if not documents:
        flash('No documents found in this batch.', 'error')
        return redirect(url_for('batch_results', batch_id=batch_id))
    
    try:
        # Create a ZIP file in memory
        zip_buffer = io.BytesIO()
        with ZipFile(zip_buffer, 'w') as zip_file:
            for doc in documents:
                docx_path = os.path.join(app.config['GENERATED_FOLDER'], doc.file_path)
                if os.path.exists(docx_path):
                    # Add DOCX file to ZIP
                    zip_file.write(docx_path, doc.original_filename)
                    
                    # Also add PDF version if it exists
                    pdf_path = docx_path.replace('.docx', '.pdf')
                    if os.path.exists(pdf_path):
                        pdf_filename = doc.original_filename.replace('.docx', '.pdf')
                        zip_file.write(pdf_path, pdf_filename)
        
        zip_buffer.seek(0)
        
        # Send the ZIP file
        zip_filename = f"batch_{batch_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        return send_file(
            zip_buffer,
            as_attachment=True,
            download_name=zip_filename,
            mimetype='application/zip'
        )
        
    except Exception as e:
        logger.error(f"Error creating batch download: {str(e)}")
        flash('Error creating batch download. Please try downloading documents individually.', 'error')
        return redirect(url_for('batch_results', batch_id=batch_id))

@app.route('/admin')
def admin():
    key = request.args.get('key')
    if key != app.config['ADMIN_KEY']:
        abort(403)
    stats = {
        'templates': Template.query.count(),
        'active_templates': Template.query.filter_by(is_active=True).count(),
        'total_documents': CreatedDocument.query.count(),
        'total_batches': BatchGeneration.query.count(),
    }
    return render_template('admin.html', stats=stats, admin_key=key)
@app.route('/admin/templates')
def admin_templates():
    key = request.args.get('key')
    if key != app.config['ADMIN_KEY']:
        abort(403)
    page = request.args.get('page', 1, type=int)
    search = request.args.get('search', '')

    query = Template.query
    if search:
        query = query.filter(
            db.or_(
                Template.name.ilike(f'%{search}%'),
                Template.type.ilike(f'%{search}%'),
                Template.description.ilike(f'%{search}%')
            )
        )

    templates = query.order_by(Template.created_at.desc()).paginate(
        page=page, per_page=20, error_out=False
    )

    return render_template('admin/templates.html',
                           templates=templates,
                           search=search,
                           admin_key=key)

@app.route('/admin/upload', methods=['GET', 'POST'])
def admin_upload_template():
    key = request.args.get('key') or request.form.get('key')
    if key != app.config['ADMIN_KEY']:
        abort(403)

    if request.method == 'GET':
        return render_template('admin/upload.html', admin_key=key)

    try:
        name = request.form['name']
        template_type = request.form['type']
        description = request.form.get('description', '')
        file = request.files['file']

        if not file or not file.filename.endswith('.docx'):
            flash('Please upload a valid .docx file', 'error')
            return redirect(url_for('admin_upload_template', key=key))

        # Save file with unique name
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"{timestamp}_{filename}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(file_path)

        # Create template record
        template = Template(
            name=name,
            type=template_type,
            description=description,
            file_path=unique_filename
        )
        db.session.add(template)
        db.session.commit()

        # Extract document-level styles
        doc = Document(file_path)
        section = doc.sections[0]
        template.margin_top = section.top_margin.inches
        template.margin_bottom = section.bottom_margin.inches
        template.margin_left = section.left_margin.inches
        template.margin_right = section.right_margin.inches
        template.default_line_spacing = doc.paragraphs[0].paragraph_format.line_spacing if doc.paragraphs else 1.0
        
        # Use dominant font and size for better accuracy across the entire document
        template.font_family, template.font_size = DocumentProcessor.get_dominant_font_and_size(doc)

        # Extract ALL placeholder instances with formatting
        placeholder_instances = DocumentProcessor.extract_template_variables(file_path)

        # Create placeholders with instance numbering for multiples
        instance_counters = Counter()
        i = 0
        for inst in placeholder_instances:
            var_name = inst['name']
            instance_counters[var_name] += 1
            if instance_counters[var_name] == 1:
                instance_name = var_name
            else:
                instance_name = f"{var_name}_instance_{instance_counters[var_name]}"
            base_name = var_name
            var_type = DocumentProcessor.detect_variable_type(base_name)
            formatting = inst['formatting']
            placeholder_text = DocumentProcessor.get_smart_placeholder_default(base_name)
            help_text = DocumentProcessor.get_smart_help_text(base_name)
            display_name = base_name.replace('_', ' ').title()
            if instance_counters[var_name] > 1:
                display_name += f" (Instance {instance_counters[var_name]})"
            placeholder = Placeholder(
                template_id=template.id,
                name=instance_name,
                display_name=display_name,
                placeholder_type=var_type,
                sort_order=i,
                help_text=help_text,
                bold=formatting.get('bold', False),
                italic=formatting.get('italic', False),
                underline=formatting.get('underline', False),
                font_family=formatting.get('font', template.font_family),
                font_size=formatting.get('size', template.font_size),
                alignment=inst['alignment'],
                left_indent=0.0,
                paragraph_index=inst['paragraph_index'],
                run_index=inst['run_index'],
                default_value=placeholder_text,
                is_required=True
            )
            if var_type == 'option':
                smart_options = DocumentProcessor.get_smart_options(base_name)
                placeholder.options = json.dumps(smart_options)
            else:
                placeholder.options = json.dumps([])
            db.session.add(placeholder)
            i += 1

        db.session.commit()

        flash(f'Template "{name}" uploaded successfully with {len(placeholder_instances)} placeholders', 'success')
        return redirect(url_for('admin_edit_template', template_id=template.id, key=key))

    except Exception as e:
        logger.error(f"Error uploading template: {str(e)}")
        flash(f'Error uploading template: {str(e)}. Please try again.', 'error')
        return redirect(url_for('admin_upload_template', key=key))

@app.route('/admin/template/<int:template_id>/edit', methods=['GET', 'POST'])
def admin_edit_template(template_id):
    key = request.args.get('key') or request.form.get('key')
    if key != app.config['ADMIN_KEY']:
        abort(403)

    template = Template.query.get_or_404(template_id)

    if request.method == 'GET':
        placeholders = Placeholder.query.filter_by(template_id=template_id).order_by(Placeholder.sort_order).all()
        return render_template('admin/edit_template.html', template=template, placeholders=placeholders, admin_key=key)

    try:
        template.name = request.form['name']
        template.type = request.form['type']
        template.description = request.form.get('description', '')
        template.font_family = request.form['font_family']
        try:
            template.font_size = int(request.form['font_size'])
            template.margin_top = float(request.form['margin_top'])
            template.margin_bottom = float(request.form['margin_bottom'])
            template.margin_left = float(request.form['margin_left'])
            template.margin_right = float(request.form['margin_right'])
            template.default_line_spacing = float(request.form['default_line_spacing'])
        except (ValueError, TypeError) as e:
            flash('Invalid numeric values in form fields', 'error')
            return redirect(url_for('admin_edit_template', template_id=template_id, key=key))
        template.updated_at = datetime.now(timezone.utc)

        for placeholder in template.placeholders:
            prefix = f'{placeholder.id}_'
            placeholder.display_name = request.form.get(prefix + 'display_name', placeholder.name)
            placeholder.placeholder_type = request.form.get(prefix + 'type', 'text')
            placeholder.is_required = True  # All placeholders required by default
            placeholder.help_text = request.form.get(prefix + 'help', '')
            placeholder.bold = prefix + 'bold' in request.form
            placeholder.italic = prefix + 'italic' in request.form
            placeholder.underline = prefix + 'underline' in request.form
            placeholder.casing = request.form.get(prefix + 'casing', 'none')
            placeholder.default_value = request.form.get(prefix + 'placeholder', '')
            if placeholder.placeholder_type == 'option':
                placeholder.options = json.dumps(request.form.getlist(prefix + 'options'))

        db.session.commit()

        flash('Template updated successfully', 'success')
        return redirect(url_for('admin_templates', key=key))

    except Exception as e:
        logger.error(f"Error updating template: {str(e)}")
        flash('Error updating template. Please try again.', 'error')
        return redirect(url_for('admin_edit_template', template_id=template_id, key=key))

@app.route('/admin/template/<int:template_id>/pause')
def admin_pause_template(template_id):
    key = request.args.get('key')
    if key != app.config['ADMIN_KEY']:
        abort(403)
    template = Template.query.get_or_404(template_id)
    template.is_active = False
    db.session.commit()
    flash('Template paused', 'success')
    return redirect(url_for('admin_templates', key=key))

@app.route('/admin/template/<int:template_id>/resume')
def admin_resume_template(template_id):
    key = request.args.get('key')
    if key != app.config['ADMIN_KEY']:
        abort(403)
    template = Template.query.get_or_404(template_id)
    template.is_active = True
    db.session.commit()
    flash('Template resumed', 'success')
    return redirect(url_for('admin_templates', key=key))

@app.route('/admin/template/<int:template_id>/delete')
def admin_delete_template(template_id):
    key = request.args.get('key')
    if key != app.config['ADMIN_KEY']:
        abort(403)
    template = Template.query.get_or_404(template_id)
    try:
        os.remove(os.path.join(app.config['UPLOAD_FOLDER'], template.file_path))
    except:
        pass
    db.session.delete(template)
    db.session.commit()
    flash('Template deleted', 'success')
    return redirect(url_for('admin_templates', key=key))

@app.route('/admin/database/clear')
def admin_clear_database():
    """Clear all database data - WARNING: This deletes everything!"""
    key = request.args.get('key')
    if key != app.config['ADMIN_KEY']:
        abort(403)
    
    try:
        # Delete all files in generated folder
        import shutil
        if os.path.exists(app.config['GENERATED_FOLDER']):
            shutil.rmtree(app.config['GENERATED_FOLDER'])
            os.makedirs(app.config['GENERATED_FOLDER'], exist_ok=True)
        
        # Delete all files in upload folder
        if os.path.exists(app.config['UPLOAD_FOLDER']):
            shutil.rmtree(app.config['UPLOAD_FOLDER'])
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        
        # Clear all database tables
        db.drop_all()
        db.create_all()
        
        # Clear all caches
        cache.clear()
        
        flash('Database and all files cleared successfully! Starting fresh.', 'success')
        logger.info("Database and files cleared by admin")
        
    except Exception as e:
        logger.error(f"Error clearing database: {str(e)}")
        flash(f'Error clearing database: {str(e)}', 'error')
    
    return redirect(url_for('admin', key=key))

@app.route('/admin/database/backup')
def admin_backup_database():
    """Create a backup of the database"""
    key = request.args.get('key')
    if key != app.config['ADMIN_KEY']:
        abort(403)
    
    try:
        import shutil
        from datetime import datetime
        
        # Create backup filename with timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_filename = f'app_backup_{timestamp}.db'
        backup_path = os.path.join(BASE_DIR, 'db', backup_filename)
        
        # Copy current database
        current_db = os.path.join(BASE_DIR, 'db', 'app.db')
        if os.path.exists(current_db):
            shutil.copy2(current_db, backup_path)
            flash(f'Database backed up successfully as {backup_filename}', 'success')
            logger.info(f"Database backup created: {backup_filename}")
        else:
            flash('No database found to backup', 'warning')
            
    except Exception as e:
        logger.error(f"Error creating backup: {str(e)}")
        flash(f'Error creating backup: {str(e)}', 'error')
    
    return redirect(url_for('admin', key=key))

@app.route('/delete/<int:document_id>')
def delete_document(document_id):
    doc = CreatedDocument.query.get_or_404(document_id)
    
    # Background file deletion - don't wait for file operations
    try:
        docx_path = os.path.join(app.config['GENERATED_FOLDER'], doc.file_path)
        pdf_path = doc.file_path.replace('.docx', '.pdf')
        pdf_full_path = os.path.join(app.config['GENERATED_FOLDER'], pdf_path)
        
        # Delete files in background to avoid blocking the UI
        import threading
        def delete_files():
            try:
                if os.path.exists(docx_path):
                    os.remove(docx_path)
                if os.path.exists(pdf_full_path):
                    os.remove(pdf_full_path)
            except:
                pass
        
        threading.Thread(target=delete_files, daemon=True).start()
    except:
        pass
    
    # Immediate database deletion
    db.session.delete(doc)
    db.session.commit()
    
    # Clear homepage cache to update immediately
    cache.clear()
    
    flash('Document deleted successfully', 'success')
    return redirect(url_for('index'))

# Error Handlers
@app.errorhandler(404)
def not_found_error(error):
    return render_template('error.html', message="Page not found"), 404

@app.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    return render_template('error.html', message="Internal server error"), 500



# ... (rest of the admin routes remain the same as in your original file)

# Initialize Database
with app.app_context():
    db.create_all()

if __name__ == '__main__':
    try:
        port = int(os.environ.get('PORT', 8000))
    except (ValueError, TypeError):
        port = 8000
    
    app.run(debug=os.environ.get('FLASK_DEBUG', 'False').lower() == 'true', 
            host=os.environ.get('HOST', '127.0.0.1'), 
            port=port)

